import sys
import os
import html
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QVBoxLayout, QHBoxLayout, QWidget, QLabel, QFrame,
    QFileDialog, QMessageBox, QSplitter, QScrollArea, QCheckBox, QPushButton, QDialog,
    QGridLayout, QGraphicsOpacityEffect
)
from PySide6.QtCore import Qt, QTimer, QThread, Signal, QPropertyAnimation, QEasingCurve, QObject
from PySide6.QtGui import QCursor, QFont, QColor, QPalette 

# --- 依赖库安全导入 ---
try:
    import chardet 
    HAS_CHARDET = True
except ImportError:
    HAS_CHARDET = False

try:
    import docx 
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import fitz 
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

from ui_components import (
    Theme, ThreeDButton, GlowingButton, ModernProgressBar, 
    AIGCGaugeWidget, AIGCPieChart, HeatmapBar, DragTextEdit, ResultBlock, StatsDashboard, DetailedHeatmapWindow,
    DeveloperConsole, HistoryWindow, EmptyStateWidget
)
from core_engine import AIGCDetectionThread, get_resource_path, check_gpu_availability, load_settings, save_settings, load_history, save_history, clear_all_history

# ================= 跨线程信号桥梁 =================
class APIMonitor(QObject):
    """安全地将后台 Flask 线程的心跳传递到主 GUI 线程"""
    heartbeat = Signal()

# ---------------------- 静默硬件嗅探线程 ----------------------
class HWScannerThread(QThread):
    finished_scan = Signal(bool, str)
    def run(self):
        has_gpu, msg = check_gpu_availability()
        self.finished_scan.emit(has_gpu, msg)

# ---------------------- 主程序窗口 ----------------------
class AIGCSentinel(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("DeepVeri - AI 生成检测系统") 
        self.resize(1300, 850)
        
        self.is_model_valid = False 
        self.model_path = ""
        self.last_results = []      
        self.detailed_heatmap_win = None 
        self._is_restoring = False 
        
        self.engine_config = load_settings()
        self.has_gpu = False
        self.gpu_name = "检测中..."
        self.is_hw_scanned = False 
        
        self.render_queue = []
        self.render_timer = QTimer(self)
        self.render_timer.timeout.connect(self._process_render_batch)
        
        # --- API 心跳状态管理 ---
        self.api_monitor = APIMonitor()
        self.api_monitor.heartbeat.connect(self._on_api_heartbeat)
        
        self.api_timeout_timer = QTimer(self)
        self.api_timeout_timer.setInterval(5000) # 5秒没收到心跳就判定断开
        self.api_timeout_timer.timeout.connect(self._on_api_timeout)
        
        Theme.CURRENT_MODE = 'dark'
        
        self.init_ui()           
        self.update_theme()      
        self.check_model_status() # 这里面会自动拉起 API 微服务
        
        self.hw_scanner = HWScannerThread()
        self.hw_scanner.finished_scan.connect(self._on_hw_scanned)
        self.hw_scanner.start()

    def _on_hw_scanned(self, has_gpu, msg):
        self.is_hw_scanned = True
        self.has_gpu = has_gpu
        self.gpu_name = msg
        self.status_text.setText(f"✅ 硬件嗅探完毕: {msg}。控制台已就绪！")
        self.status_text.setStyleSheet(f"color: {Theme.ACCENT_GREEN.name()}; font-weight: bold; padding: 0 4px;")

    def init_ui(self):
        central = QWidget()
        central.setObjectName("centralWidget") 
        self.setCentralWidget(central)
        layout = QVBoxLayout(central)
        layout.setContentsMargins(20, 20, 20, 15) 
        layout.setSpacing(15) 

        header = QHBoxLayout()
        title_box = QVBoxLayout()
        
        self.title_lbl = QLabel("DeepVeri")
        self.title_lbl.setStyleSheet("font-size: 24px; font-weight: 900; letter-spacing: 1.5px;") 
        
        self.sub_lbl = QLabel("深度学习文本 AI 生成检测平台") 
        self.sub_lbl.setStyleSheet("font-size: 11px; font-weight: bold; letter-spacing: 1px; color: #2D79FF;")
        
        title_box.addWidget(self.title_lbl)
        title_box.addWidget(self.sub_lbl)
        header.addLayout(title_box)
        header.addSpacing(40)
        
        self.btn_history = GlowingButton("🕒  历史", variant="secondary", parent=self)
        self.btn_history.setFixedWidth(85)
        self.btn_history.clicked.connect(self.open_history)
        header.addWidget(self.btn_history)
        
        self.btn_console = GlowingButton("⚙️  控制台", variant="secondary", parent=self)
        self.btn_console.setFixedWidth(90)
        self.btn_console.clicked.connect(self.open_console)
        header.addWidget(self.btn_console)
        header.addSpacing(15)
        
        self.btn_clear = GlowingButton("🗑️  清空内容", variant="danger", parent=self)
        self.btn_clear.setFixedWidth(105)
        self.btn_clear.clicked.connect(self.clear_content)
        header.addWidget(self.btn_clear)
        header.addStretch() 
        
        self.btn_import = GlowingButton("📂  导入文档", variant="secondary", parent=self)
        self.btn_import.setFixedWidth(115)
        self.btn_import.clicked.connect(self.import_file)
        
        self.btn_merge = GlowingButton("✂️  合并排版", variant="secondary", parent=self)
        self.btn_merge.setFixedWidth(115)
        self.btn_merge.clicked.connect(self.merge_all_lines)
        
        self.btn_detect = GlowingButton("⚡  开始深度检测", variant="primary", parent=self)
        self.btn_detect.setFixedWidth(160)
        self.btn_detect.clicked.connect(self.run_detection)
        
        header.addWidget(self.btn_import)
        header.addSpacing(12)
        header.addWidget(self.btn_merge)
        header.addSpacing(12)
        header.addWidget(self.btn_detect)
        layout.addLayout(header)

        splitter = QSplitter(Qt.Horizontal)
        splitter.setHandleWidth(20)
        
        self.card_input = QFrame()
        in_layout = QVBoxLayout(self.card_input)
        in_header = QHBoxLayout()
        
        self.label_input = QLabel("📝  原文输入 (支持 .txt / .docx / .pdf 拖入)")
        self.label_input.setStyleSheet("font-weight: bold; margin-bottom: 5px;")
        
        self.lbl_char_count = QLabel("字数: 0")
        in_header.addWidget(self.label_input)
        in_header.addStretch()
        in_header.addWidget(self.lbl_char_count)
        
        self.input_edit = DragTextEdit()
        self.input_edit.file_dropped.connect(self.handle_file_content)
        self.input_edit.textChanged.connect(self.update_char_count) 
        
        in_layout.addLayout(in_header)
        in_layout.addWidget(self.input_edit)

        self.card_output = QFrame() 
        out_outer_layout = QHBoxLayout(self.card_output)
        out_outer_layout.setContentsMargins(0, 10, 5, 10)
        
        self.right_stack_layout = QGridLayout()
        self.right_stack_layout.setContentsMargins(0, 0, 0, 0)
        
        self.content_container = QWidget()
        res_main_layout = QVBoxLayout(self.content_container)
        res_main_layout.setContentsMargins(0, 0, 0, 0)
        
        self.dashboard = StatsDashboard()
        res_main_layout.addWidget(self.dashboard)
        
        self.gauge = self.dashboard.gauge
        self.pie_chart = self.dashboard.pie_chart
        self.token_counter = self.dashboard.token_counter
        
        ctrl_bar = QHBoxLayout()
        self.label_output = QLabel("🔍  逐段 AI 生成检测")
        self.label_output.setStyleSheet("font-weight: bold; font-size: 14px;")
        
        self.chk_only_high_risk = QCheckBox("只显示高风险内容 (>60%)")
        self.chk_only_high_risk.setCursor(Qt.PointingHandCursor)
        self.chk_only_high_risk.stateChanged.connect(self.apply_filter) 
        
        ctrl_bar.addWidget(self.label_output)
        ctrl_bar.addStretch()
        ctrl_bar.addWidget(self.chk_only_high_risk)
        ctrl_bar.addSpacing(10)
        res_main_layout.addLayout(ctrl_bar)
        
        self.result_scroll = QScrollArea()
        self.result_scroll.setWidgetResizable(True)
        self.result_scroll.setFrameShape(QFrame.NoFrame)
        self.result_container = QWidget()
        self.result_layout = QVBoxLayout(self.result_container)
        self.result_layout.setAlignment(Qt.AlignTop)
        self.result_layout.setSpacing(10)
        self.result_scroll.setWidget(self.result_container)
        res_main_layout.addWidget(self.result_scroll)
        
        self.right_stack_layout.addWidget(self.content_container, 0, 0)
        
        self.empty_container = EmptyStateWidget()
        self.right_stack_layout.addWidget(self.empty_container, 0, 0)
        
        self.content_container.hide()
        self.empty_container.show()
        
        self.heatmap = HeatmapBar()
        self.heatmap.clicked_section.connect(self.scroll_to_section) 
        self.heatmap.double_clicked.connect(self.show_detailed_heatmap)

        out_outer_layout.addLayout(self.right_stack_layout)
        out_outer_layout.addWidget(self.heatmap)

        splitter.addWidget(self.card_input)
        splitter.addWidget(self.card_output)
        splitter.setSizes([600, 500]) 
        layout.addWidget(splitter, stretch=1)

        # ------------------ 底部状态栏 ------------------
        status_bar = QFrame()
        status_bar.setFixedHeight(24)
        sb_layout = QHBoxLayout(status_bar)
        sb_layout.setContentsMargins(0, 0, 0, 0)
        
        self.status_icon = QLabel("●")
        self.status_text = QLabel("初始化...")
        
        self.btn_refresh = QPushButton("🔄  刷新")
        self.btn_refresh.setFixedSize(70, 20) 
        self.btn_refresh.clicked.connect(self.manual_refresh_model)
        
        sb_layout.addWidget(self.status_icon)
        sb_layout.addWidget(self.status_text)
        sb_layout.addSpacing(5)
        sb_layout.addWidget(self.btn_refresh)
        sb_layout.addStretch()
        
       # --- 【核心修改】：Vercel 风格极简指示灯 (初始为离线状态) ---
        self.api_status_lbl = QLabel("<span style='font-size: 16px; vertical-align: middle;'>•</span> 节点离线")
        self.offline_api_css = "color: #EF4444; font-weight: bold; font-size: 11px; margin-right: 15px;" 
        self.api_status_lbl.setStyleSheet(self.offline_api_css)
        
        # --- 【核心新增】：为指示灯创建专属的透明度特效和脉冲动画 ---
        self.api_pulse_eff = QGraphicsOpacityEffect(self.api_status_lbl)
        self.api_status_lbl.setGraphicsEffect(self.api_pulse_eff)
        
        self.api_pulse_anim = QPropertyAnimation(self.api_pulse_eff, b"opacity")
        self.api_pulse_anim.setDuration(400) # 闪烁总耗时 0.4 秒
        self.api_pulse_anim.setStartValue(1.0)
        self.api_pulse_anim.setKeyValueAt(0.4, 0.4) # 中间变暗，模拟心跳
        self.api_pulse_anim.setEndValue(1.0)
        self.api_pulse_anim.setEasingCurve(QEasingCurve.InOutQuad)
        
        sb_layout.addWidget(self.api_status_lbl)
        
        self.label_device = QLabel("")
        self.label_device.setStyleSheet("color: #666; font-size: 11px; margin-right: 10px;")
        sb_layout.addWidget(self.label_device)
        
        self.progress_bar = ModernProgressBar()
        self.progress_bar.setFixedWidth(300)
        sb_layout.addWidget(self.progress_bar)
        layout.addWidget(status_bar)

    # ------------------ API 节点交互逻辑 ------------------
    def _on_api_heartbeat(self):
        """收到外界请求，变为极简绿点，并触发心跳闪烁动效"""
        active_css = "color: #10B981; font-weight: bold; font-size: 11px; margin-right: 15px;"
        
        # 只有在状态真正改变时才更新 UI，减少重绘开销
        if self.api_status_lbl.text() != "<span style='font-size: 16px; vertical-align: middle;'>•</span> 节点活跃":
            self.api_status_lbl.setText("<span style='font-size: 16px; vertical-align: middle;'>•</span> 节点活跃")
            self.api_status_lbl.setStyleSheet(active_css)
        
        # 触发脉冲闪烁动画 (如果上一次没播完，强行重播)
        if self.api_pulse_anim.state() == QPropertyAnimation.Running:
            self.api_pulse_anim.stop()
        self.api_pulse_anim.start()
        
        self.api_timeout_timer.start()

    def _on_api_timeout(self):
        """超过倒计时没收到请求，瞬间变回离线红点"""
        self.api_status_lbl.setText("<span style='font-size: 16px; vertical-align: middle;'>•</span> 节点离线")
        self.api_status_lbl.setStyleSheet(self.offline_api_css)
        
        self.api_timeout_timer.stop()
        self.api_pulse_eff.setOpacity(1.0) # 确保透明度恢复正常

    # ------------------ 历史记录面板交互 ------------------
    def open_history(self):
        history_data = load_history()
        self.history_dlg = HistoryWindow(history_data, self)
        self.history_dlg.request_restore.connect(self.restore_from_history)
        self.history_dlg.request_clear.connect(self.clear_history_data)
        self.history_dlg.exec()

    def clear_history_data(self):
        reply = QMessageBox.question(
            self, 
            "确认清空", 
            "确定要清空所有检测历史记录吗？\n此操作不可恢复。", 
            QMessageBox.Yes | QMessageBox.No
        )
        if reply == QMessageBox.Yes:
            clear_all_history()
            if hasattr(self, 'history_dlg') and self.history_dlg:
                self.history_dlg.clear_list()
            QMessageBox.information(self, "成功", "历史记录已彻底清空。")

    def restore_from_history(self, record):
        if self.detailed_heatmap_win:
            self.detailed_heatmap_win.close()
            
        self.clear_content(show_empty=False) 
        self.show_content_panel()
        
        text = record.get("original_text", "")
        html_content = f"<div style='line-height: 1.6;'>{html.escape(text).replace(chr(10), '<br>')}</div>"
        self.input_edit.setHtml(html_content)
        
        res = {
            "total_ai_rate": record.get("total_ai_rate", 0),
            "total_tokens": record.get("total_tokens", 0),
            "paragraphs": record.get("paragraphs", [])
        }
        
        self._is_restoring = True
        self.process_results(res)
        self._is_restoring = False
        
        self.status_text.setText(f"✅ 已成功恢复历史记录：{record.get('timestamp')}")
        self.status_text.setStyleSheet(f"color: {Theme.ACCENT_GREEN.name()}; font-weight: bold; padding: 0 4px;")

    def open_console(self):
        if not self.is_hw_scanned:
            QMessageBox.information(self, "稍等", "系统正在初始化并静默扫描底层显卡硬件池，请等待几秒钟再打开控制台。")
            return
            
        dlg = DeveloperConsole(
            current_config=self.engine_config, 
            has_gpu=self.has_gpu, 
            gpu_name=self.gpu_name, 
            parent=self
        )
        
        if dlg.exec() == QDialog.Accepted:
            self.engine_config.update(dlg.config)
            save_settings(self.engine_config) 
            
            mode = "纯 CPU 算力接管" if self.engine_config['force_cpu'] else "智能硬件加速"
            QMessageBox.information(
                self, 
                "底层引擎参数已更新", 
                f"新参数（切分阈值 {self.engine_config['max_chunk_size']} 等）已永久保存。\n"
                f"将在下一次「开始深度检测」时以 {mode} 生效！"
            )
            
            if self.engine_config['force_cpu']:
                self.update_device_ui("🐢 预载：用户强制切断硬件加速 (CPU)", False)
            else:
                self.update_device_ui(f"🚀 预载：准备调用 {self.gpu_name}", True)

    def manual_refresh_model(self):
        self.btn_refresh.setEnabled(False) 
        bg = Theme.get('bg_main')
        
        self.status_text.setText("正在扫描本地模型...")
        self.status_text.setStyleSheet(f"color: {Theme.ACCENT_YELLOW.name()}; background-color: {bg}; font-weight: bold; padding: 0 4px;")
        self.status_text.repaint()
        QApplication.processEvents() 
        
        QThread.msleep(300) 
        self.check_model_status()
        
        if self.is_model_valid:
            QMessageBox.information(self, "状态更新", "成功检测到本地模型！")
        else:
            QMessageBox.warning(self, "状态更新", "仍然未检测到完整模型。")
        self.btn_refresh.setEnabled(True)

    def check_model_status(self):
        target_dir = get_resource_path("AIGC_Model")
        if not os.path.exists(target_dir):
            self.set_model_invalid("未找到 'AIGC_Model' 文件夹")
            return
            
        try:
            files = os.listdir(target_dir)
            bg = Theme.get('bg_main')
            has_config = "config.json" in files
            has_weights = "pytorch_model.bin" in files or "model.safetensors" in files
            
            if has_config and has_weights:
                self.is_model_valid = True
                self.model_path = target_dir
                self.status_icon.setStyleSheet(f"color: {Theme.ACCENT_GREEN.name()}; font-size: 14px;")
                self.status_text.setText("本地引擎已就绪")
                self.status_text.setStyleSheet(f"color: {Theme.ACCENT_GREEN.name()}; background-color: {bg}; font-weight: bold; padding: 0 4px;")
                
                # --- 核心改动：在确信模型可用后，把跨线程回调挂载给 API 微服务启动函数 ---
                if not hasattr(self, '_api_started'):
                    try:
                        from api_server import start_api_server
                        start_api_server(
                            self.model_path, 
                            self.engine_config, 
                            port=5005, 
                            notify_callback=self.api_monitor.heartbeat.emit  # 将 Signal 指针传给 Flask 线程
                        )
                        self._api_started = True
                    except Exception as e:
                        print(f"微服务静默启动失败: {e}")
            else:
                self.set_model_invalid("缺失核心权重文件")
        except Exception as e:
            self.set_model_invalid(f"读取异常: {str(e)}")

    def set_model_invalid(self, reason):
        self.is_model_valid = False
        bg = Theme.get('bg_main')
        self.status_icon.setStyleSheet(f"color: {Theme.ACCENT_RED.name()}; font-size: 14px;")
        self.status_text.setText(f"⚠️ 无法检测: {reason}")
        self.status_text.setStyleSheet(f"color: {Theme.ACCENT_RED.name()}; background-color: {bg}; font-weight: bold; padding: 0 4px;")

    def update_device_ui(self, msg, is_gpu):
        self.label_device.setText(msg)
        color = Theme.ACCENT_GREEN.name() if is_gpu else Theme.ACCENT_YELLOW.name()
        if "强制切断" in msg or "错误" in msg: 
            color = Theme.ACCENT_RED.name()
        self.label_device.setStyleSheet(f"color: {color}; font-weight: bold; font-size: 11px; margin-right: 15px;")

    def show_detailed_heatmap(self):
        if not hasattr(self, 'last_results') or not self.last_results:
            return
        if self.detailed_heatmap_win and self.detailed_heatmap_win.isVisible():
            self.detailed_heatmap_win.activateWindow()
        else:
            self.detailed_heatmap_win = DetailedHeatmapWindow(self.last_results, self)
            self.detailed_heatmap_win.request_scroll.connect(self.scroll_to_section)
            self.detailed_heatmap_win.show()

    def run_detection(self):
        if hasattr(self, 'work_thread') and self.work_thread.isRunning():
            self.work_thread.stop()
            self.btn_detect.setEnabled(False)
            self.btn_detect.setText("正在终止...")
            self.status_text.setText("正在安全切断检测流...")
            return

        if not self.is_model_valid:
            QMessageBox.critical(self, "无法运行", "请先确保模型文件夹已正确放置。")
            return
            
        text = self.input_edit.toPlainText().strip()
        if not text: 
            self.btn_detect.setText("⚠️ 内容为空")
            QTimer.singleShot(1500, lambda: self.btn_detect.setText("⚡  开始深度检测")) 
            return
            
        self.btn_detect.setVariant("danger")
        self.btn_detect.setText("⏹️ 终止检测")
        self.show_content_panel()
        
        self.render_timer.stop()
        self.render_queue = []
        
        while self.result_layout.count() > 0:
            item = self.result_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
                
        self.gauge.setValue(0)
        self.progress_bar.setValue(0)
        
        self.work_thread = AIGCDetectionThread(text, self.model_path, config=self.engine_config)
        self.work_thread.status_signal.connect(lambda s: self.status_text.setText(s))
        self.work_thread.progress_signal.connect(self.progress_bar.setValue)
        self.work_thread.result_signal.connect(self.process_results)
        self.work_thread.device_signal.connect(self.update_device_ui)
        self.work_thread.finished.connect(self._on_thread_finished)
        self.work_thread.start()

    def _on_thread_finished(self):
        self.btn_detect.setEnabled(True)
        self.btn_detect.setVariant("primary")
        self.btn_detect.setText("⚡  开始深度检测")
        self.progress_bar.setValue(100)

    def process_results(self, res):
        if "error" in res:
            QMessageBox.critical(self, "检测中断", res["error"])
            return
        
        self.last_results = res.get("paragraphs", [])
        
        if not self._is_restoring:
            save_history(
                original_text=self.input_edit.toPlainText().strip(),
                total_ai_rate=res["total_ai_rate"],
                total_tokens=res.get("total_tokens", 0),
                paragraphs=self.last_results
            )
        
        self.gauge.setValue(res["total_ai_rate"])
        self.token_counter.set_data(res.get("total_tokens", 0))
        self.heatmap.set_data(self.last_results) 
        
        counts = [0, 0, 0] 
        for p in self.last_results:
            if p.get("is_ignored"): continue
            rate = p["ai_rate"]
            if rate < 30: counts[0] += 1
            elif rate < 60: counts[1] += 1
            else: counts[2] += 1
            
        self.pie_chart.set_data(counts)

        self.render_queue = list(enumerate(self.last_results))
        QTimer.singleShot(200, lambda: self.render_timer.start(17)) 
        
        if hasattr(self, 'detailed_heatmap_win') and self.detailed_heatmap_win and self.detailed_heatmap_win.isVisible():
            self.detailed_heatmap_win.close()
            self.show_detailed_heatmap() 

    def _process_render_batch(self):
        if not self.render_queue:
            self.render_timer.stop()
            self.result_layout.addStretch() 
            self.apply_filter() 
            return

        self.result_container.setUpdatesEnabled(False)
        
        for _ in range(2):
            if not self.render_queue: break
            idx, p = self.render_queue.pop(0)
            use_anim = (idx < 10)
            
            block = ResultBlock(
                index=idx, 
                content=p["content"], 
                ai_rate=p["ai_rate"], 
                tokens=p.get("tokens", 0),
                is_ignored=p.get("is_ignored", False), 
                use_animation=use_anim
            )
            block.request_scroll.connect(self.handle_block_resize) 
            block.request_highlight.connect(self.highlight_source_text) 
            block.expanded.connect(self.on_block_expanded) 
            
            self.result_layout.addWidget(block)
            
            if self.chk_only_high_risk.isChecked() and p["ai_rate"] <= 60:
                block.hide()

        self.result_container.setUpdatesEnabled(True)
        self.result_container.update()

    def on_block_expanded(self, expanded_index):
        for i in range(self.result_layout.count()):
            item = self.result_layout.itemAt(i)
            widget = item.widget()
            if widget and isinstance(widget, ResultBlock):
                if widget.index != expanded_index and widget.is_expanded:
                    widget.set_expanded(False)

    def highlight_source_text(self, content):
        self.input_edit.highlight_paragraph(content)

    def apply_filter(self):
        show_high = self.chk_only_high_risk.isChecked()
        for i in range(self.result_layout.count()):
            item = self.result_layout.itemAt(i)
            widget = item.widget()
            if widget and isinstance(widget, ResultBlock):
                if show_high: 
                    widget.setVisible(widget.ai_rate > 60)
                else: 
                    widget.show()

    def handle_block_resize(self):
        self.result_container.adjustSize()

    def scroll_to_section(self, index):
        target = None
        for i in range(self.result_layout.count()):
            widget = self.result_layout.itemAt(i).widget()
            if widget and isinstance(widget, ResultBlock) and widget.index == index:
                target = widget
                break
        
        if target:
            if target.isHidden():
                self.chk_only_high_risk.setChecked(False) 
                QApplication.processEvents() 
                
            if not target.is_expanded:
                target.toggle_expand() 
                
            QApplication.processEvents() 
            self.result_scroll.verticalScrollBar().setValue(target.pos().y())
            self.result_scroll.horizontalScrollBar().setValue(0)
            self.highlight_source_text(target.content) 

    def clear_content(self, *args, show_empty=True):
        self.render_timer.stop() 
        self.render_queue = []
        self.input_edit.clear()
        self.lbl_char_count.setText("字数: 0") 
        self.last_results = []
        
        if show_empty:
            self.show_empty_panel()
        
        while self.result_layout.count() > 0:
            item = self.result_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
                
        self.gauge.setValue(0)
        self.progress_bar.setValue(0)
        self.heatmap.set_data([])
        self.pie_chart.set_data([0, 0, 0])
        self.token_counter.set_data(0)  
        
        if self.detailed_heatmap_win:
            self.detailed_heatmap_win.close() 

    def show_content_panel(self):
        if not self.content_container.isVisible():
            if hasattr(self.empty_container, 'pause_breathing'):
                self.empty_container.pause_breathing()

            self.empty_eff = QGraphicsOpacityEffect()
            self.empty_container.setGraphicsEffect(self.empty_eff)

            self.content_eff = QGraphicsOpacityEffect()
            self.content_eff.setOpacity(0.0)
            self.content_container.setGraphicsEffect(self.content_eff)

            self.content_container.show()

            self.fade_out_empty = QPropertyAnimation(self.empty_eff, b"opacity")
            self.fade_out_empty.setDuration(300)
            self.fade_out_empty.setStartValue(1.0)
            self.fade_out_empty.setEndValue(0.0)
            self.fade_out_empty.setEasingCurve(QEasingCurve.OutQuad)

            self.fade_in_content = QPropertyAnimation(self.content_eff, b"opacity")
            self.fade_in_content.setDuration(400)
            self.fade_in_content.setStartValue(0.0)
            self.fade_in_content.setEndValue(1.0)
            self.fade_in_content.setEasingCurve(QEasingCurve.OutQuad)

            def _on_fade_out():
                self.empty_container.hide()
                self.empty_container.setGraphicsEffect(None) 

            def _on_fade_in():
                self.content_container.setGraphicsEffect(None) 

            self.fade_out_empty.finished.connect(_on_fade_out)
            self.fade_in_content.finished.connect(_on_fade_in)

            self.fade_out_empty.start()
            self.fade_in_content.start()

    def show_empty_panel(self):
        if self.content_container.isVisible():
            self.content_eff = QGraphicsOpacityEffect()
            self.content_container.setGraphicsEffect(self.content_eff)

            self.empty_eff = QGraphicsOpacityEffect()
            self.empty_eff.setOpacity(0.0)
            self.empty_container.setGraphicsEffect(self.empty_eff)

            self.empty_container.show()

            self.fade_out_content = QPropertyAnimation(self.content_eff, b"opacity")
            self.fade_out_content.setDuration(300)
            self.fade_out_content.setStartValue(1.0)
            self.fade_out_content.setEndValue(0.0)
            self.fade_out_content.setEasingCurve(QEasingCurve.OutQuad)

            self.fade_in_empty = QPropertyAnimation(self.empty_eff, b"opacity")
            self.fade_in_empty.setDuration(400)
            self.fade_in_empty.setStartValue(0.0)
            self.fade_in_empty.setEndValue(1.0)
            self.fade_in_empty.setEasingCurve(QEasingCurve.OutQuad)

            def _on_fade_out():
                self.content_container.hide()
                self.content_container.setGraphicsEffect(None)

            def _on_fade_in():
                self.empty_container.setGraphicsEffect(None)
                if hasattr(self.empty_container, 'resume_breathing'):
                    self.empty_container.resume_breathing()

            self.fade_out_content.finished.connect(_on_fade_out)
            self.fade_in_empty.finished.connect(_on_fade_in)

            self.fade_out_content.start()
            self.fade_in_empty.start()

    def merge_all_lines(self):
        text = self.input_edit.toPlainText()
        if not text.strip(): 
            return
            
        import re
        text = re.sub(r'([\u4e00-\u9fa5])\s*\n\s*([\u4e00-\u9fa5])', r'\1\2', text)
        text = text.replace('\n', ' ')
        
        html_content = f"<div style='line-height: 1.6;'>{html.escape(text).replace(chr(10), '<br>')}</div>"
        self.input_edit.setHtml(html_content)
        self.status_text.setText("✅ 已合并排版结构，并保持舒适行距")

    def handle_file_content(self, path):
        QApplication.setOverrideCursor(QCursor(Qt.WaitCursor)) 
        try:
            ext = os.path.splitext(path)[1].lower()
            content = ""
            
            if ext == '.txt':
                with open(path, 'rb') as f:
                    raw = f.read()
                encoding = chardet.detect(raw)['encoding'] if HAS_CHARDET else 'utf-8'
                content = raw.decode(encoding, errors='ignore')
            elif ext == '.docx':
                if not HAS_DOCX:
                    raise Exception("缺失 python-docx 库")
                doc = docx.Document(path)
                text_parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        unique_cells = []
                        seen_cells = set()
                        for cell in row.cells:
                            if cell not in seen_cells:
                                unique_cells.append(cell)
                                seen_cells.add(cell)
                        row_text_list = []
                        for cell in unique_cells:
                            txt = cell.text.strip()
                            if txt:
                                row_text_list.append(txt)
                        if row_text_list:
                            text_parts.append(" | ".join(row_text_list))
                content = "\n\n".join(text_parts)
            elif ext == '.pdf':
                if not HAS_PDF:
                    raise Exception("缺失 PyMuPDF 库")
                import re
                doc = fitz.open(path)
                text_parts = []
                for page in doc:
                    blocks = page.get_text("blocks")
                    for b in blocks:
                        if b[6] == 0:  
                            text = b[4].strip()
                            if text:
                                text = re.sub(r'([\u4e00-\u9fa5])\s*\n\s*([\u4e00-\u9fa5])', r'\1\2', text)
                                text = text.replace('\n', ' ')
                                text_parts.append(text)
                content = "\n\n".join(text_parts)
            
            html_content = f"<div style='line-height: 1.6;'>{html.escape(content).replace(chr(10), '<br>')}</div>"
            self.input_edit.setHtml(html_content)
            self.status_text.setText(f"已成功加载并提取全部内容: {os.path.basename(path)}")
            
        except Exception as e:
            QMessageBox.critical(self, "读取失败", str(e))
        finally:
            QApplication.restoreOverrideCursor() 

    def import_file(self):
        path, _ = QFileDialog.getOpenFileName(self, "打开文档", "", "支持的文件 (*.txt *.docx *.pdf)")
        if path:
            self.handle_file_content(path)

    def update_char_count(self):
        count = len(self.input_edit.toPlainText())
        self.lbl_char_count.setText(f"字数: {count:,}")

    def update_theme(self):
        t = Theme.COLORS['dark']
        
        palette = QApplication.palette()
        palette.setColor(QPalette.Window, QColor(t['bg_main']))
        palette.setColor(QPalette.WindowText, QColor(t['text_main']))
        palette.setColor(QPalette.Base, QColor(t['input_bg']))
        palette.setColor(QPalette.AlternateBase, QColor(t['bg_card']))
        palette.setColor(QPalette.Text, QColor(t['text_main']))
        palette.setColor(QPalette.Button, QColor(t['bg_card']))
        palette.setColor(QPalette.ButtonText, QColor(t['text_main']))
        palette.setColor(QPalette.Highlight, Theme.ACCENT_BLUE)
        QApplication.setPalette(palette)
        
        self.setStyleSheet(f"""
            QMainWindow, #centralWidget {{ background-color: {t['bg_main']}; }}
            QSplitter::handle {{ background: transparent; }}
            QScrollArea {{ background: transparent; border: none; }}
            QCheckBox {{ color: {t['text_sub']}; font-family: 'Microsoft YaHei'; font-weight: bold; }}
            QCheckBox::indicator {{ width: 18px; height: 18px; border-radius: 6px; border: 1px solid {t['border']}; }}
            QCheckBox::indicator:checked {{ background-color: #3B82F6; border-color: #3B82F6; image: url(data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHZpZXdCb3g9IjAgMCAyNCAyNCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSJ3aGl0ZSIgc3Ryb2tlLXdpZHRoPSIzIiBzdHJva2UtbGluZWNhcD0icm91bmQiIHN0cm9rZS1saW5lam9pbj0icm91bmQiPjxwb2x5bGluZSBwb2ludHM9IjIwIDYgOSAxNyA0IDEyIi8+PC9zdmc+); }}
            {Theme.SCROLLBAR_CSS}
        """)
        
        self.input_edit.setStyleSheet(f"""
            QTextEdit {{ 
                background-color: {t['input_bg']}; 
                color: {t['text_main']}; 
                border: 1px solid rgba(255, 255, 255, 0.04); 
                border-radius: 16px; 
                padding: 24px 22px; 
                font-size: 11.5pt; 
                selection-background-color: #3B82F6;
                selection-color: white;
            }}
            QTextEdit:focus {{ 
                border: 1px solid rgba(59, 130, 246, 0.5); 
                background-color: {QColor(t['input_bg']).lighter(102).name()};
            }}
        """)
        
        self.card_input.setStyleSheet(f"QFrame {{ background-color: {t['bg_card']}; border: 1px solid {t['border']}; border-radius: 16px; }}")
        self.card_output.setStyleSheet(f"QFrame {{ background-color: {t['bg_card']}; border: 1px solid {t['border']}; border-radius: 16px; }}")
        self.card_input.setGraphicsEffect(Theme.shadow(35))
        self.card_output.setGraphicsEffect(Theme.shadow(35))
        
        if hasattr(self, 'dashboard'):
            self.dashboard.update_style()
            
        self.btn_refresh.setStyleSheet("""
            QPushButton { 
                background-color: rgba(255,255,255,0.05); 
                color: #9CA3AF; 
                border-radius: 10px; 
                border: 1px solid rgba(255,255,255,0.05); 
                font-size: 11px; 
                font-weight: bold; 
                padding: 0 12px; 
            } 
            QPushButton:hover { 
                background-color: #3B82F6; 
                color: white; 
                border: 1px solid #3B82F6; 
            }
        """)

if __name__ == "__main__":
    os.environ["QT_AUTO_SCREEN_SCALE_FACTOR"] = "1"
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    font = QFont("Microsoft YaHei", 10)
    font.setStyleStrategy(QFont.PreferAntialias)
    app.setFont(font)
    
    window = AIGCSentinel()
    window.show()
    sys.exit(app.exec())